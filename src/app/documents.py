#!/usr/bin/env python3
"""
documents.py
Purpose: Document generation from chat conversations
Status: [ACTIVE]
Created: 2025-09-13
Author: writeian

Handles generating structured documents (notes, outlines) from chat discussions
for educational scaffolding without academic dishonesty concerns.
"""

from flask import (
    Blueprint,
    request,
    redirect,
    url_for,
    flash,
    current_app,
    send_file,
    make_response,
)
from datetime import datetime
import io
import tempfile
import os
from typing import Any, List
from src.app import db
from src.models import Chat, Message
from src.app.access_control import get_current_user, require_chat_access

documents = Blueprint("documents", __name__)


@documents.route("/chat/<int:chat_id>/generate", methods=["POST"])
@require_chat_access
def generate_from_chat(chat_id: int) -> Any:
    """Generate a structured document from chat conversation."""
    try:
        chat_obj = Chat.query.get_or_404(chat_id)
        user = get_current_user()
        
        # Get all messages for this chat
        messages = Message.query.filter_by(chat_id=chat_obj.id).order_by(Message.timestamp).all()
        
        # Check minimum message threshold
        if len(messages) < 5:
            flash(f"Need at least 5 messages to generate a document. Current: {len(messages)} messages.", "warning")
            return redirect(url_for("chat.view_chat", chat_id=chat_obj.id))
        
        # Get document type and format from request
        doc_type = request.form.get("doc_type", "notes")
        format_type = request.form.get("format", "text")
        
        # Generate document content based on chat messages
        document_content = generate_document_content(messages, chat_obj, doc_type)
        
        # Create downloadable file based on format
        if format_type == "docx":
            return create_docx_download(document_content, chat_obj, doc_type)
        else:
            # Default to plain text/markdown
            return create_text_download(document_content, chat_obj, doc_type)
            
    except Exception as e:
        current_app.logger.error(f"Error generating document for chat {chat_id}: {e}")
        flash("Failed to generate document. Please try again.", "error")
        return redirect(url_for("chat.view_chat", chat_id=chat_obj.id))


def generate_document_content(messages: List[Message], chat_obj: Chat, doc_type: str) -> str:
    """Generate structured document content from chat messages."""
    from src.utils.openai_utils import call_anthropic_api
    from src.app.room.utils.room_utils import infer_template_type_from_room
    
    # Prepare chat content for AI analysis
    chat_content = []
    for msg in messages:
        role_prefix = "User" if msg.role == "user" else "AI Assistant"
        if msg.user and msg.role == "user":
            role_prefix = f"{msg.user.display_name}"
        chat_content.append(f"{role_prefix}: {msg.content}")
    
    chat_text = "\n\n".join(chat_content)
    
    # Determine template type for context
    template_type = infer_template_type_from_room(chat_obj.room) or "general"
    
    # Create AI prompt based on document type
    prompts = {
        "notes": f"""
Based on the following chat discussion from a {template_type} learning session, create structured notes that help students organize their thinking. DO NOT write full content for them - provide frameworks, key points, and questions they need to answer.

Chat Discussion:
{chat_text}

Create structured notes with:
1. Key questions discussed (with space for student answers)
2. Main arguments identified (with space for evidence)
3. Important concepts mentioned (with space for analysis)
4. Next steps needed (action items for students)

Format as clear, organized notes that scaffold learning without doing the work for students.
""",
        "outline": f"""
Based on the following chat discussion from a {template_type} learning session, create a document outline that helps students structure their work. Provide the framework but leave the actual content for students to write.

Chat Discussion:
{chat_text}

Create an outline with:
- Clear section headings based on discussion topics
- Bullet points for key areas discussed
- [Student fills in] placeholders for actual content
- Guiding questions for each section
- Space for student research and analysis

This should be a roadmap for student work, not completed content.
""",
        "summary": f"""
Based on the following chat discussion, create an organized summary that captures the key insights and structures the thinking. Focus on frameworks and organization, not completed content.

Chat Discussion:
{chat_text}

Create a helpful organizational structure that students can build upon.
"""
    }
    
    prompt = prompts.get(doc_type, prompts["summary"])
    
    try:
        # Generate document content using AI
        content, _ = call_anthropic_api(
            [{"role": "user", "content": prompt}],
            system_prompt="You are an expert educator who helps students organize their thinking without doing their work for them. Provide scaffolding, frameworks, and questions - not completed assignments.",
            max_tokens=1500
        )
        return content
    except Exception as e:
        current_app.logger.error(f"Error generating document content: {e}")
        return f"""# Discussion Notes

Generated from chat: {chat_obj.title}
Date: {datetime.now().strftime('%Y-%m-%d')}
Room: {chat_obj.room.name}

## Key Discussion Points

[Content generation failed - please try again]

## Messages in Discussion: {len(messages)}

This document would normally contain organized notes from your chat discussion.
"""


def create_text_download(content: str, chat_obj: Chat, doc_type: str) -> Any:
    """Create a downloadable text file."""
    filename = f"{doc_type}_{chat_obj.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt"
    
    # Create response with text content
    response = make_response(content)
    response.headers['Content-Type'] = 'text/plain; charset=utf-8'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response


def create_docx_download(content: str, chat_obj: Chat, doc_type: str) -> Any:
    """Create a downloadable Word document."""
    try:
        # Try to import python-docx
        from docx import Document
        from docx.shared import Inches
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f"{doc_type.title()}: {chat_obj.title}", 0)
        
        # Add metadata
        doc.add_paragraph(f"Generated from AI Collab chat discussion")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Room: {chat_obj.room.name}")
        doc.add_paragraph(f"Learning Mode: {chat_obj.mode}")
        doc.add_paragraph("")  # Empty line
        
        # Add content (parse markdown-style formatting)
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                # Main heading
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                # Sub heading
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                # Sub-sub heading
                doc.add_heading(line[4:], 3)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                # Numbered list
                doc.add_paragraph(line[3:], style='List Number')
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        filename = f"{doc_type}_{chat_obj.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        def cleanup_temp_file():
            try:
                os.unlink(temp_file.name)
            except:
                pass
        
        # Create response
        response = send_file(
            temp_file.name,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Schedule cleanup
        import atexit
        atexit.register(cleanup_temp_file)
        
        return response
        
    except ImportError:
        # Fallback to text if python-docx not available
        current_app.logger.warning("python-docx not available, falling back to text export")
        return create_text_download(content, chat_obj, doc_type)
    except Exception as e:
        current_app.logger.error(f"Error creating DOCX file: {e}")
        return create_text_download(content, chat_obj, doc_type)


def get_available_document_types(message_count: int, template_type: str = None) -> List[dict]:
    """Get available document types based on message count and room template."""
    types = []
    
    if message_count >= 5:
        types.append({
            "key": "notes",
            "label": "📝 Discussion Notes",
            "description": "Organized notes with key points and questions"
        })
    
    if message_count >= 10:
        types.append({
            "key": "outline", 
            "label": "📋 Document Outline",
            "description": "Structured framework for writing"
        })
        
        # Template-specific options
        if template_type == "academic-essay":
            types.append({
                "key": "essay_framework",
                "label": "📄 Essay Framework", 
                "description": "Essay structure with thesis development"
            })
        elif template_type == "business-hub":
            types.append({
                "key": "business_summary",
                "label": "💼 Business Summary",
                "description": "Executive summary and action items"
            })
    
    return types
